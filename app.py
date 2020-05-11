
import csv
import os

from flask import Flask, session,send_file,render_template,redirect,flash, request, url_for,jsonify, Response
from flask_session import Session
from flask_login import LoginManager, UserMixin, login_required, login_user, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
# Database
# from config import Config
# from config import Config
from sqlalchemy import event, create_engine, or_, and_
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import scoped_session, sessionmaker
from flask_migrate import Migrate
from flask_login import LoginManager
# from flask_uploads import UploadSet, IMAGES, configure_uploads

# from sqlalchemy.exc import IntegrityError
from sqlalchemy.event import listen
from werkzeug import secure_filename
import time
import datetime
import json
import barcode, random
from barcode.writer import ImageWriter

from models import Users,Students, Levels, Majors, Subjects, Company

# from flask_wtf import FlaskForm
# from wtforms import StringField
# from wtforms.validators import DataRequired

# Import Python Docx to Generate Word File
from docx import Document
from docx.shared import Inches

document = Document()

from flask import Blueprint
from flask_paginate import Pagination, get_page_parameter, get_page_args
#
mod = Blueprint('student', __name__)

app = Flask(__name__)

# UPLOAD_FOLDER = '/files/'
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
#
# configure session to use filesystem (instead of signed cookies)
# app.config["SESSION_FILE_DIR"] = gettempdir()
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"
#
app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'
Session(app)

# DB Connection Postgress

# app.config.from_object(os.environ['APP_SETTINGS'])
# app.config.from_object(Config)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_DATABASE_URI'] = "postgresql://postgres:123456789@localhost:5432/student_barcode"
db = SQLAlchemy(app)
migrate = Migrate(app, db)

# POSTGRES = {
#     'user': 'postgres',
#     'pw': 'password',
#     'db': 'student_barcode',
#     'host': 'localhost',
#     'port': '5432',
# }
# app.config['DEBUG'] = True
# app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://%(user)s:%(pw)s@%(host)s:%(port)s/%(db)s' % POSTGRES
# db.init_app(app)

# engine = create_engine('postgresql://scott:tiger@localhost/mydatabase')
# migrate = Migrate(app, db)

# images = UploadSet('images', IMAGES)
# configure_uploads(app, images)

def login_required(f):
    """
    Decorate routes to require login.

    http://flask.pocoo.org/docs/0.12/patterns/viewdecorators/
    """
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("user_id") is None:
            return redirect("/login")
        return f(*args, **kwargs)
    return decorated_function



# second solution
@event.listens_for(Users.__table__, 'after_create')
def insert_initial_values(*args, **kwargs):
    db.session.add(Users(name='Ahmed Eid',username='admin',email='ahmedmohamedeid97@gmail.com',password=generate_password_hash("admin"),is_super_user=True, is_admin=True))
    # db.session.add(Users(name='medium'))
    # db.session.add(Users(name='high'))
    db.session.commit()

# event.listen(Users.__table__, 'after_create', insert_initial_values)


from sqlalchemy.ext.declarative import DeclarativeMeta

class AlchemyEncoder(json.JSONEncoder):

    def default(self, obj):
        if isinstance(obj.__class__, DeclarativeMeta):
            # an SQLAlchemy class
            fields = {}
            for field in [x for x in dir(obj) if not x.startswith('_') and x != 'metadata']:
                data = obj.__getattribute__(field)
                try:
                    json.dumps(data) # this will fail on non-encodable values, like other classes
                    fields[field] = data
                except TypeError:
                    fields[field] = None
            # a json-encodable dict
            return fields

        return json.JSONEncoder.default(self, obj)


   # User Register Page #
@app.route("/signup", methods=['GET','POST'])
def signup():
    # session.clear()

    if request.method == 'POST':
        # name = request.form.get("full_name")
        username = request.form.get("name")
        email = request.form.get("email")
        password = generate_password_hash(request.form.get("password"))
        print (password)
        # check_email = Users.query(db.exists().where(Users.email==email)).scalar()
        check_email = db.session.query(db.exists().where(Users.email == email)).scalar()
        print(check_email)
        if not check_email:
            user = Users(name=username,username=username,email=email,password=password)
            print(user)
            db.session.add(user)
            db.session.commit()

            # Get Username to store Session
            session['userid'] = user.id
            session['username'] = user.username
            # redirect to Home
            return redirect("/")
        else:
            msg = "Theis User is already Exist!"
            print(msg)
            return render_template("signup.html", message=msg)
    else:
        message = "Invalid Register"
        print(message)
        return render_template("signup.html", message=message)

# print("email", email)
# print("pass", password)
# valid_email = False
# valid_pass = False
# if email == None:
#     flash("البريد الالكتروني فارغ...")
#     valid_email = False
# elif len(email) == 8:
#     flash("البريد الالكتروني اقل من 8 احرف.")
#     valid_email = False
# else:
#     valid_email = True
# if password == None:
#     flash("كلمة المرور فارغة ...")
#     valid_pass = False
# elif len(password) == 8:
#     flash(" كلمة المرور اقل من 8 احرف.")
#     valid_pass = False
# else:
#     valid_pass = True
# print(valid_email)
# print(valid_pass)

@app.route('/<path:path>/wizard', methods=('GET', 'POST'))
def wizard_setting(path):

    if path == "company" and request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        phone = request.form.get("phone")
        page_url = request.form.get("page_url")
        notes = request.form.get("notes")
        nu_of_days_for_lock = request.form.get("nu_of_days_for_lock")

        lock_at = datetime.datetime.now() + datetime.timedelta(days=int(nu_of_days_for_lock))

        status = True

        file = request.files['logo']
        image_name = file.filename

        res = Company(name=name, email=email, phone=phone, page_url=page_url,notes=notes,nu_of_days_for_lock=nu_of_days_for_lock,lock_at=lock_at,image_name=image_name,image=file.read())
        db.session.add(res)
        db.session.merge(res)
        db.session.commit()
        flash("تم الاضافة بنجاح")
        return redirect(url_for('wizard_setting', path="users"))
    if path == "company":

        return render_template('wizard_setting.html', company=path)

    elif path == "users" and request.method == "POST" :
        username = request.form.get("name")
        email = request.form.get("email")
        password = generate_password_hash(request.form.get("password"))
        is_admin = request.form.get("is_admin")
        is_user = request.form.get("is_user")
        status = request.form.get("status")

        is_admin = True if is_admin == "1" else False

        is_user = True if is_user == "1" else False

        status = True if status == "1" else False

        check_count = db.session.query(Users).count()
        is_super_user = True if check_count == 0 else False
        if is_super_user:
            is_admin = True
        if is_admin:
            is_user = True

        check_email = db.session.query(db.exists().where(Users.email == email)).scalar()
        if not check_email:
            user = Users(name=username,username=username,email=email,password=password,is_super_user=is_super_user, is_admin=is_admin, is_user=is_user, status=status)
            db.session.add(user)
            db.session.commit()
        flash("تم أضافة المستخدم بنجاح!")
        return redirect("/login")
    elif path == "users":

        company = Company.query.all()
        if company:
            return render_template('wizard_setting.html', companys_user=company, users='users')
        else:
            return render_template('wizard_setting.html', company="company")


@app.route('/login', methods=('GET', 'POST'))
def login():

    if request.method == "POST":
        email = request.form.get("username")
        password = request.form.get("password")
        hash_password = generate_password_hash(password)

        result = Users.query.filter_by(email=email).first()

        if result:
            company = Company.query.get(result.company_id)
            current_date = datetime.datetime.now()
            if current_date > company.lock_at:
                all_users = Users.query.filter(Users.is_super_user != True).all()
                for user in all_users:
                    user.status = False
                    db.session.merge(user)
                db.session.commit()

            if check_password_hash(result.password, password) and result.status == True and result.is_user == True:
                session["user_id"] = result.id
                session["user_name"] = result.name
                session["is_super_user"] = result.is_super_user
                session["is_admin"] = result.is_admin
                session["company_id"] = result.company_id
                result.last_login = datetime.datetime.now()
                print(result.last_login)
                db.session.merge(result)
                db.session.commit()
                return redirect('/')
            else:
                flash('Invalid email or password or your account not approived')
                return redirect("/login")
        else:
            flash("Please Check your Email or password")
            return render_template("login.html")
        return render_template("login.html")
    else:
        company = Company.query.all()
        if company:
            return render_template('login.html')
        else:
            return redirect(url_for('wizard_setting', path="company"))

@app.route("/logout")
@login_required
def logout():
    # logout_user()
    session.clear()
    return redirect('/login')


@app.route('/setting')
@login_required
def setting():
    majors =  Majors.query.all()
    levels =  Levels.query.all()
    company = Company.query.all()
    if session["is_super_user"]:
        users  =  Users.query.filter(Users.id != session["user_id"]).all()
    else:
        users  =  Users.query.filter(Users.id != session["user_id"], or_(Users.is_super_user == False,Users.is_super_user == None)).all()


    # rows = {
    #     levels: levels[0],
    # }
    return render_template('setting.html', company=company,users=users, levels=levels, majors=majors)

@app.route('/addcompany', methods=["POST"])
@login_required
def company():
    if request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        phone = request.form.get("phone")
        page_url = request.form.get("page_url")
        notes = request.form.get("notes")
        nu_of_days_for_lock = request.form.get("nu_of_days_for_lock")

        lock_at = datetime.datetime.now() + datetime.timedelta(days=int(nu_of_days_for_lock))

        res = Company(name=name, email=email, phone=phone, page_url=page_url,notes=notes,nu_of_days_for_lock=nu_of_days_for_lock,lock_at=lock_at)
        db.session.add(res)
        db.session.merge(res)
        db.session.commit()
        flash("تم الاضافة بنجاح")
        return redirect("/setting")
    else:
        return redirect("/logout")


@app.route('/adduser', methods=["POST"])
@login_required
def users():
    if request.method == 'POST':
        username = request.form.get("name")
        email = request.form.get("email")
        password = generate_password_hash(request.form.get("password"))
        is_admin = request.form.get("is_admin")
        is_user = request.form.get("is_user")
        status = request.form.get("status")

        is_admin = True if is_admin == "1" else False

        is_user = True if is_user == "1" else False

        status = True if status == "1" else False

        check_count = db.session.query(Users).count()
        is_super_user = True if check_count == 0 else False
        if is_super_user:
            is_admin = True
        if is_admin:
            is_user = True

        check_email = db.session.query(db.exists().where(Users.email == email)).scalar()
        if not check_email:
            user = Users(name=username,username=username,email=email,password=password,is_super_user=is_super_user, is_admin=is_admin, is_user=is_user, status=status)
            db.session.add(user)
            db.session.commit()
        flash("تم أضافة المستخدم بنجاح!")
        return redirect("/setting")
    else:
        flash("خطأ")
        return redirect("/setting")


@app.route('/major', methods=["POST"])
@login_required
def major():
    name = request.form.get("name")
    levels = request.form.getlist("levels")

    print(levels)
    major = Majors(name=name )
    db.session.add(major)
    db.session.commit()
    current_major = Majors.query.filter_by(name=name).first()
    print(current_major.id)
    # for level in levels:
    #     # level_id = Levels.query.filter_by(name=level)
    #     add = Majors.levels(level_id=level, major_id=current_major.id)
    #     print(add)
    # db.session.commit()
    # print("list :",level_list)
    #
    # # major.levels.append(levels)
    # # level = Levels(name=name,notes=notes)
    # print(major)
    # # db.session.add(major)
    # db.session.commit()
    return redirect("/setting")

@app.route('/level', methods=["POST"])
@login_required
def level():
    name = request.form.get("name")
    notes = request.form.get("notes")

    level = Levels(name=name,notes=notes)
    print(level)
    db.session.add(level)
    db.session.commit()
    return redirect("/setting")

from io import BytesIO
# edit For All tables
@app.route('/event/<int:id>/logo')
def event_logo(id):
    company = Company.query.get_or_404(id)
    return app.response_class(company.image, mimetype='application/octet-stream')

@app.route('/<path:path>/edit/<int:id>', methods=["POST","GET"])
@login_required
def updateRecords(path, id):
    if path == "company" and request.method == 'POST':
        if id:
            name = request.form.get("name")
            email = request.form.get("email")
            phone = request.form.get("phone")
            page_url = request.form.get("page_url")
            notes = request.form.get("notes")
            nu_of_days_for_lock = request.form.get("nu_of_days_for_lock")
            status = request.form.get('status')
            status = True if status == "1" else False

            file = request.files['logo']
            image_name = file.filename
            print(image_name)
            print(file)
            company = Company.query.get(id)
            if nu_of_days_for_lock !=None and nu_of_days_for_lock != company.nu_of_days_for_lock:
                company.lock_at = company.create_at + datetime.timedelta(days=int(nu_of_days_for_lock))

            company.name = name
            company.email = email
            company.phone = phone
            company.page_url = page_url
            company.notes = notes
            company.update_at = datetime.datetime.now()
            if nu_of_days_for_lock !=None and nu_of_days_for_lock != company.nu_of_days_for_lock :
                company.nu_of_days_for_lock = nu_of_days_for_lock
            if status !=  None and status != company.status:
                company.status = status
            if file:
                company.image_name = image_name
                company.image = file.read()

            db.session.merge(company)
            db.session.commit()
            return redirect("/setting")
    elif path == "company":
        company = Company.query.get(id)
        return render_template('setting_edite.html', company=company)
    elif path == "users" and request.method == 'POST':
        name = request.form.get("name")
        username = request.form.get("name")
        email = request.form.get("email")
        company_id = request.form.get("company_id")
        # password = generate_password_hash(request.form.get("password"))
        is_admin = request.form.get("is_admin")
        is_user = request.form.get("is_user")
        status = request.form.get("status")

        is_admin = True if is_admin == "1" else False

        is_user = True if is_user == "1" else False

        status = True if status == "1" else False

        users = Users.query.get(id)

        if is_admin:
            is_user = True
        # print(company_id)
        users.company_id = company_id
        users.name = name
        users.email = email
        users.is_admin = is_admin
        users.is_user = is_user
        users.status = status
        users.update_at = datetime.datetime.now()

        db.session.merge(users)
        db.session.commit()
        return redirect("/setting")
    elif path == 'users':
        users = Users.query.get(id)
        companys_user = Company.query.all()
        return render_template('setting_edite.html', users=users, companys_user=companys_user )
    elif path == "majors" and request.method == 'POST':
        name = request.form.get("name")
        # levels = request.form.getlist("levels")
        notes = request.form.get("notes")
        major = Majors.query.get(id)
        major.name = name
        major.notes = notes
        major.update_at = datetime.datetime.now()

        db.session.merge(major)
        db.session.commit()
        return redirect("/setting")
    elif path == 'majors':
        major = Majors.query.get(id)
        return render_template('setting_edite.html', major=major)

    elif path == "levels" and request.method == 'POST':
        name = request.form.get("name")
        # levels = request.form.getlist("levels")
        notes = request.form.get("notes")
        level = Levels.query.get(id)
        level.name = name
        level.notes = notes
        level.update_at = datetime.datetime.now()

        db.session.merge(level)
        db.session.commit()
        return redirect("/setting")
    elif path == 'levels':
        level = Levels.query.get(id)
        return render_template('setting_edite.html', level=level)

def get_users(offset=0, per_page=10):
    return users[offset: offset + per_page]

@app.route('/')
@app.route('/index')
@login_required
def index():

    majors =  Majors.query.all()
    levels =  Levels.query.all()
    print(majors)
    # page, per_page, offset = get_page_args(page_parameter='page',per_page_parameter='per_page')
    # Print(page)
    page = request.args.get(get_page_parameter(), type=int, default=1)
    print(page)

    # users = student.find(...)
    # rows = db.execute("SELECT * FROM student").fetchall()
    rows = Students.query.all()
    # print(rows)
    pagination = Pagination(page=page, css_framework="bootstrap4",total=len(rows), record_name='rows')
    print(pagination)
    return render_template("index.html", pagination=pagination ,students=rows,majors=majors, levels=levels)

#
#
#
@app.route("/upload", methods=["GET","POST"])
@login_required
def student_upload_data():
    """Upload Student Data."""

    if request.method == "POST":
        major = request.form.get("major")
        level = request.form.get("level")
        file = request.files['student_file']
        file.save(secure_filename(file.filename))

        file = open(file.filename)
        students = csv.reader(file)

        for ui_code,name in students:
            students = Students(name=name,un_id=ui_code,major_id=major,level_id=level)
            db.session.add(students)
        db.session.commit()
        print(file)
    return redirect(url_for("index"))

    # else:
    #     return render_template("register.html")
#

@app.route('/students')
@login_required
def all_student():
    majors =  Majors.query.all()
    levels =  Levels.query.all()
    page = request.args.get(get_page_parameter(), type=int, default=1)
    # print(page)
    rows = Students.query.all()
    rowss = {
        "students":rows,
    }
    major_level_dec = []
    for row in rows:
        # print(row.major_id.)
        major =  Majors.query.get(row.major_id)
        level =  Levels.query.get(row.level_id)
        major_level_dec.append({
        "student_id": row.id,
        "major":major.name,
        "level": level.name
        })
    rowss["data"] = major_level_dec
    # print("================",rowss)
        # print(major_level_dec)
    pagination = Pagination(page=page, css_framework="bootstrap4",total=len(rows), record_name='rows')
    # print(pagination)

    return render_template("student.html", pagination=pagination ,rows=rowss,major_level=major_level_dec,majors=majors,levels=levels )

# @app.route('/user/<username>')
# def profile(username):
#     ...
#
# @app.route('/<int:year>/<int:month>/<title>')
# def article(year, month, title):

@app.route('/student/<int:id>')
@login_required
def student(id):
    row = Students.query.get(id)
    major_id =  Majors.query.get(row.major_id)
    level_id =  Levels.query.get(row.level_id)

    return render_template("student_data.html",student=row,major=major_id,level=level_id )



# # c = YourAlchemyClass()
# # print (json.dumps(c, cls=AlchemyEncoder))
#
# @app.route('/search', methods=["GET","POST"])
# def search():
#     if request.method == 'GET':
#         text = request.args['searchText']
#         # search_data = request.form.get('search_data')
#         # print(search_data)
#         search = "%{}%".format(text)
#         result = Students.query.filter(Students.name.like(search)).all()
#
#         print(json.dumps(result, cls=AlchemyEncoder))
#         return json.dumps({"success": True, 'data':json.dumps(result, cls=AlchemyEncoder)})
#         # return render_template("student.html", pagination=pagination ,rows=rowss,major_level=major_level_dec,majors=majors,levels=levels )
#
from sqlalchemy import and_, or_, not_
@app.route('/search', methods=["GET","POST"])
def search():
    if request.method == 'POST':
        text = request.form.get('search_data')
        major_search = request.form.get('major_filter')
        level_search = request.form.get('level_filter')

        if isinstance(text, int):
            result = Students.query.filter(Students.un_id == text ).all()
        elif isinstance(text, str):
            search = "%{}%".format(text)
            result = Students.query.filter(Students.name.like(search) ).all()

        if major_search and level_search:
            result = Students.query.filter(Students.major_id == major_search,Students.level_id == level_search).all()
        majors =  Majors.query.all()
        levels =  Levels.query.all()
        page = request.args.get(get_page_parameter(), type=int, default=1)
        # print(page)

        rowss = {
            "students":result,
        }
        major_level_dec = []
        for row in result:
            # print(row.major_id.)
            major =  Majors.query.get(row.major_id)
            level =  Levels.query.get(row.level_id)
            major_level_dec.append({
            "student_id": row.id,
            "major":major.name,
            "level": level.name
            })
        rowss["data"] = major_level_dec
        # print("================",rowss)
            # print(major_level_dec)
        pagination = Pagination(page=page, css_framework="bootstrap4",total=len(result), record_name='rows')
        # print(pagination)

        return render_template("student.html", pagination=pagination ,rows=rowss,major_level=major_level_dec,majors=majors,levels=levels )


@app.route("/delete/<path:path>/<int:id>")
@login_required
def deleterecord(path,id):
    if path == 'levels':
        level = Levels.query.filter_by(id=id).first()
        merge_level = db.session.merge(level)
        db.session.delete(merge_level)
        db.session.commit()
        flash('Successfully Delete')
        return redirect(url_for('setting'))
    elif path == 'majors':
        major = Majors.query.filter_by(id=id).first()
        merge_major = db.session.merge(major)
        db.session.delete(merge_major)
        db.session.commit()
        flash('Successfully Delete')
        return redirect(url_for('setting'))
    elif path == 'users':
        user = Users.query.filter_by(id=id).first()
        merge_user = db.session.merge(user)
        db.session.delete(merge_user)
        db.session.commit()
        flash('Successfully Delete')
        return redirect(url_for('setting'))
    elif path == 'company':
        company = Company.query.filter_by(id=id).first()
        merge_company = db.session.merge(company)
        db.session.delete(merge_company)
        db.session.commit()
        flash('Successfully Delete')
        return redirect(url_for('setting'))

# Generate Word File
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH


@login_required
@app.route("/generate_word_file/<int:id>")
def generate_word_file(id):
    # style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    row = Students.query.get(id)
    major_id =  Majors.query.get(row.major_id)
    level_id =  Levels.query.get(row.level_id)

    head_title = document.add_heading(row.name, 0)
    head_title.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # document.add_heading('Heading, level 1', level=1)
    nu_id = document.add_paragraph('الكود الجامعي : {}'.format(row.un_id))
    nu_id.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    major = document.add_paragraph('الشعبة : {}'.format(major_id.name))
    major.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    level = document.add_paragraph('الفرقة : {}'.format(level_id.name))
    level.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subjects = document.add_paragraph('المواد الدراسية : ')
    subjects.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_picture('icon.png', width=Inches(1.25))
    #
    # table = document.add_table(row =3, cols= 3)
    # table.direction = WD_TABLE_DIRECTION.RTL

    document.add_page_break()
    path = document.save('demo-{}.docx'.format(row.un_id))
    print(path)
    return redirect(url_for('all_student'))
    # return send_file(path, as_attachment=True, attachment_filename="demo.docx")

#
# @app.route("/generate_png", methods=["GET", "POST"])
# def generate_png():
#
#     if request.method == "POST":
#         # num = random.randrange(1,00)
#         rows = db.execute("SELECT * FROM student").fetchall()
#         if rows:
#             for row in rows:
#                 num = "00{}".format(row[2])
#                 name = row[1]
#                 print (num)
#                 EAN = barcode.get_barcode_class('ean8')
#                 ean = EAN(u"{}".format(num),  writer=ImageWriter())
#                 fullname = ean.save(u'{}-{}'.format(num,name), text=name)
#     return redirect(url_for("index"))
#
@app.route("/generate_svg", methods=["GET", "POST"])
def generate_svg():
    if request.method == "POST":
        # num = random.randrange(1,00)
        path = ""
        rows = Students.query.filter(Students.barcode_path == None).all()
        # rows = db.execute("SELECT * FROM student WHERE barcode_path IS NULL;").fetchall()
        print(rows)
        barcode_path = "static/barcode_path"
        try:
            os.makedirs(barcode_path)

        except OSError:
            print ("Creation of the directory %s failed" % barcode_path)
        else:
            print ("Successfully created the directory %s " % barcode_path)
        os.chdir(barcode_path)
        if rows:
            for row in rows:

                num = "00{}".format(row.un_id)
                name = row.name
                print (num)
                EAN = barcode.get_barcode_class('ean8')
                ean = EAN(u"{}".format(num))


                filenum = random.randrange(1,1000)
                fullname = ean.save(u'{}-{}'.format(num,filenum), text=name)
                path = "barcode_path/{}-{}.svg".format(num,filenum)
                st_img_path = Students.query.get(row.id)
                st_img_path.barcode_path = path
                print(st_img_path.barcode_path)
                print(time.strftime('%A %B, %d %Y %H:%M:%S'))
                st_img_path.update_at = time.strftime('%A %B, %d %Y %H:%M:%S')
                print(fullname)
                db.session.merge(st_img_path)
                db.session.commit()
    return redirect(url_for("all_student"))
#
#
@app.route("/delete_all", methods=["POST", "GET"])
def delete_all():
    if request.method == "POST":
        # db.execute("DELETE FROM student")
        Students.query.delete()
        db.session.commit()
    return redirect(url_for("index"))


# import os
# @app.route("/print_image", methods=['GET'])
# def print_image():
#     return os.startfile("icon.png", "print")


if __name__ == '__main__':
    # app.config.update(SECRET_KEY=os.urandom(24))
    # app.secret_key = 'super secret key'
    # app.secret_key = b'_5#y2L"F4Q8z\n\xec]/'

    app.config['SESSION_TYPE'] = 'filesystem'

    sess.init_app(app)




    app.debug = True
    app.run()
